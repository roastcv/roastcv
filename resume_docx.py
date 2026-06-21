"""
resume_docx.py — DOCX resume and cover letter generator.
Uses python-docx to create Word-compatible .docx files.
ATS-friendly: single column, no tables, no text boxes.

FIX (this revision):
- Section parsing now keys sections by their canonical name directly
  (e.g. "skills") instead of the first word of the raw heading text
  (e.g. "technical" for "TECHNICAL SKILLS"). The previous approach silently
  dropped any section whose heading wasn't a single word, because the
  lookup in _get_section() never matched the stored key. The new
  _SECTION_DEFS list maps each regex pattern straight to its canonical
  key, so parsing and lookup can never disagree.
- generate_resume_docx() no longer crashes when humanized_text is empty
  (e.g. the Humanizer agent failed and an empty string was passed in).
  Name extraction now falls back safely to "" instead of indexing into
  an empty list.
"""

import re
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color helpers ────────────────────────────────────────────
def _rgb(hex_str: str) -> RGBColor:
    hex_str = hex_str.lstrip("#")
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


DARK_BLUE  = _rgb("#1a2e4a")
ACCENT     = _rgb("#2563eb")
MID_GRAY   = _rgb("#64748b")
BLACK      = _rgb("#0f172a")


# ── Paragraph border (used as section divider) ───────────────
def _add_bottom_border(paragraph, color_hex: str = "2563eb"):
    """Adds a bottom border to a paragraph — acts as a visual divider."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_spacing(paragraph, before: int = 0, after: int = 0):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _format_link(value: str, domain: str, path_prefix: str = "") -> str:
    """
    Normalizes a link field that might already be a full URL, a bare domain
    path, or just a username — avoids double-prefixing
    (e.g. "linkedin.com/in/linkedin.com/in/x").
    """
    value = value.strip()
    if value.startswith("http") or domain in value:
        return value
    return f"{path_prefix}{value.lstrip('/')}"


# ── Contact line helpers ─────────────────────────────────────
def _contact_parts(resume_data: dict) -> list:
    contact = resume_data.get("contact", {})
    parts = []
    if contact.get("email"):
        parts.append(contact["email"])
    if contact.get("phone"):
        parts.append(contact["phone"])
    if contact.get("location"):
        parts.append(contact["location"])
    if contact.get("linkedin"):
        parts.append(_format_link(contact["linkedin"], "linkedin.com", "linkedin.com/in/"))
    if contact.get("github"):
        parts.append(_format_link(contact["github"], "github.com", "github.com/"))
    if contact.get("portfolio"):
        url = contact["portfolio"]
        parts.append(url if url.startswith("http") else url.lstrip("/"))
    return parts


# ── Section parser ────────────────────────────────────────────
# FIX: each entry maps directly to the canonical key used later for lookup,
# instead of deriving the key from the first word of the matched heading
# (e.g. "TECHNICAL SKILLS" used to become key "technical", which never
# matched any lookup variant tried in _get_section() — the section just
# silently disappeared from the output document).
_SECTION_DEFS = [
    ("summary",        r"^(summary|profile|objective)"),
    ("skills",         r"^(skills|technical skills|core competencies|key skills)"),
    ("experience",     r"^(experience|work experience|employment|professional experience)"),
    ("projects",       r"^(projects|key projects|personal projects)"),
    ("education",      r"^(education|academic background|qualifications)"),
    ("certifications", r"^(certifications?|licenses?|courses?)"),
    ("achievements",   r"^(achievements?|accomplishments?|awards?)"),
]

_SECTION_ORDER = [
    ("summary",        "PROFESSIONAL SUMMARY"),
    ("skills",         "SKILLS"),
    ("experience",     "EXPERIENCE"),
    ("projects",       "PROJECTS"),
    ("education",      "EDUCATION"),
    ("certifications", "CERTIFICATIONS"),
    ("achievements",   "ACHIEVEMENTS"),
]


def _match_section_key(line: str):
    """Returns the canonical section key if `line` looks like a section
    heading, else None."""
    lowered = line.lower()
    for key, pattern in _SECTION_DEFS:
        if re.match(pattern, lowered):
            return key
    return None


def _parse_sections(text: str) -> dict:
    sections = {}
    current = "header"
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            lines.append("")
            continue
        matched_key = _match_section_key(stripped)
        if matched_key:
            sections[current] = lines
            current = matched_key
            lines = []
        else:
            lines.append(stripped)
    sections[current] = lines
    return sections


def _get_section(sections: dict, key: str) -> list:
    # FIX: sections are now stored under their canonical key directly, so a
    # plain lookup is enough — no more guessing variants like
    # "technical " + key or "work " + key.
    lines = sections.get(key, [])
    return [l for l in lines if l.strip()]


def _looks_like_heading(line: str) -> bool:
    words = line.split()
    if len(words) > 8:
        return False
    cap_words = sum(1 for w in words if w and w[0].isupper())
    return cap_words / max(len(words), 1) > 0.6


def _extract_name_from_text(text: str) -> str:
    """FIX: safe fallback — first non-empty line is usually the name.
    Returns "" instead of raising IndexError when text is empty."""
    for line in text.splitlines():
        if line.strip():
            return line.strip()
    return ""


# ── Main generator ───────────────────────────────────────────
def generate_resume_docx(
    humanized_text: str,
    resume_data: dict,
    output_path: str,
) -> str:
    """
    Generate an ATS-friendly DOCX resume.

    Args:
        humanized_text: Plain text resume from humanizer agent
        resume_data:    Structured dict from resume_reader agent
        output_path:    Where to save the .docx file
    Returns:
        output_path
    """
    doc = Document()

    # ── Page margins ────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Name ────────────────────────────────────────────────
    # FIX: was humanized_text.splitlines()[0] — crashed with IndexError when
    # humanized_text was empty (e.g. the Humanizer agent failed upstream).
    name = resume_data.get("name", "") or _extract_name_from_text(humanized_text)
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, before=0, after=40)
        run = p.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = DARK_BLUE

    # ── Contact line ─────────────────────────────────────────
    parts = _contact_parts(resume_data)
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, before=0, after=60)
        run = p.add_run("  |  ".join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GRAY

    # ── Sections ─────────────────────────────────────────────
    sections = _parse_sections(humanized_text)

    for key, heading in _SECTION_ORDER:
        lines = _get_section(sections, key)
        if not lines:
            continue

        # Section heading with bottom border
        p = doc.add_paragraph()
        _set_spacing(p, before=120, after=40)
        _add_bottom_border(p)
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT

        # Section content
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith(("- ", "• ", "* ", "· ")):
                # Bullet point
                text = line[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                _set_spacing(p, before=0, after=20)
                run = p.add_run(text)
                run.font.size = Pt(9.5)
                run.font.color.rgb = BLACK

            elif len(line) < 60 and (line.isupper() or _looks_like_heading(line)):
                # Job title / company name
                p = doc.add_paragraph()
                _set_spacing(p, before=60, after=20)
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = BLACK

            elif re.search(r"\b(20\d{2}|present|current)\b", line.lower()):
                # Date / meta line
                p = doc.add_paragraph()
                _set_spacing(p, before=0, after=20)
                run = p.add_run(line)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = MID_GRAY

            else:
                # Regular body text
                p = doc.add_paragraph()
                _set_spacing(p, before=0, after=30)
                run = p.add_run(line)
                run.font.size = Pt(9.5)
                run.font.color.rgb = BLACK

    doc.save(output_path)
    return output_path


# ── Cover Letter DOCX ─────────────────────────────────────────
def generate_cover_letter_docx(
    cover_letter_text: str,
    resume_data: dict,
    company_name: str = "",
    output_path: str = "cover_letter.docx",
) -> str:
    """
    Generate a professional DOCX cover letter.

    Args:
        cover_letter_text: Plain text cover letter from cover_letter agent
        resume_data:       Structured dict from resume_reader agent
        company_name:      Target company name
        output_path:       Where to save the .docx file
    Returns:
        output_path
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Name ────────────────────────────────────────────────
    name = resume_data.get("name", "")
    if name:
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=40)
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = DARK_BLUE

    # ── Contact ──────────────────────────────────────────────
    parts = _contact_parts(resume_data)
    if parts:
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=40)
        run = p.add_run("  |  ".join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GRAY

    # Divider
    p = doc.add_paragraph()
    _add_bottom_border(p)
    _set_spacing(p, before=0, after=120)

    # ── Date ─────────────────────────────────────────────────
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=80)
    run = p.add_run(date.today().strftime("%B %d, %Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GRAY

    # ── Company ──────────────────────────────────────────────
    if company_name:
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=120)
        run = p.add_run(f"Hiring Team, {company_name}")
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK

    # ── Body paragraphs ──────────────────────────────────────
    for para in cover_letter_text.strip().split("\n\n"):
        para = para.strip().replace("\n", " ")
        if not para:
            continue
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=140)
        run = p.add_run(para)
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK

    doc.save(output_path)
    return output_path